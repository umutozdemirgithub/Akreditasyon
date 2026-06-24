from __future__ import annotations

import json
import uuid
from datetime import datetime, timedelta
from typing import Any
from io import BytesIO

from docx import Document

from .db import get_conn, now_iso, row_to_dict, rows_to_dicts, transaction
from .repositories import (
    APPROVED,
    COMPLETED,
    EDITOR_ROLE,
    READY,
    REVISION,
    SUBMITTED,
    STATUS_OPTIONS,
    assert_program_operation_permission,
    get_section,
    list_evidence,
    list_sections,
    list_tables,
    log_activity,
    quality_for_section,
    update_section,
    get_user,
    get_program,
    program_tenant_id,
)

PUKO_FIELDS = ["planla", "uygula", "kontrol", "onlem"]
PRO_QUALITY_TARGET = 98
PRO_TARGET_LABEL = "9.8+"


def _safe_json_loads(value: Any, fallback: Any) -> Any:
    try:
        parsed = json.loads(str(value or ""))
        return parsed if parsed is not None else fallback
    except Exception:
        return fallback


def _days_until(deadline: str | None) -> int | None:
    value = str(deadline or "").strip()
    if not value:
        return None
    for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%d/%m/%Y"):
        try:
            target = datetime.strptime(value[:10], fmt)
            now = datetime.now()
            return (target.date() - now.date()).days
        except Exception:
            continue
    return None


def _assigned_owner(program_id: str, section_key: str) -> dict[str, Any]:
    with get_conn() as conn:
        rows = conn.execute(
            """SELECT pu.username, pu.role, pu.assigned_sections, COALESCE(u.full_name, pu.username) AS full_name
               FROM program_users pu
               LEFT JOIN users u ON u.username=pu.username
               WHERE pu.program_id=? AND pu.is_active=1 AND COALESCE(pu.deleted_at,'')=''
               ORDER BY CASE WHEN pu.role=? THEN 0 ELSE 1 END, COALESCE(u.full_name, pu.username)""",
            (program_id, EDITOR_ROLE),
        ).fetchall()
    for row in rows:
        raw = str(row["assigned_sections"] if "assigned_sections" in row.keys() else "" or "").strip()
        keys = {part.strip() for part in raw.split(",") if part.strip()}
        if not keys or section_key in keys:
            return {"username": row["username"], "full_name": row["full_name"], "role": row["role"]}
    return {"username": "", "full_name": "Atanmamış", "role": ""}


def _suggested_evidence(section: dict[str, Any], quality: dict[str, Any]) -> list[str]:
    title = f"{section.get('section_key','')} {section.get('section_title','')} {section.get('main_title','')}".lower()
    suggestions: list[str] = []
    if any(word in title for word in ["amaç", "hedef", "misyon", "vizyon"]):
        suggestions.extend(["Program amaçları onay tutanağı", "Paydaş görüş formları", "Amaç-hedef güncelleme toplantı kaydı"])
    if any(word in title for word in ["öğretim", "ders", "müfredat", "program çıkt"]):
        suggestions.extend(["Ders bilgi paketleri", "Program çıktısı-ders matrisi", "Müfredat güncelleme kurul kararı"])
    if any(word in title for word in ["öğrenci", "danışman", "mezun"]):
        suggestions.extend(["Öğrenci danışmanlık kayıtları", "Mezun izleme anketi", "Öğrenci memnuniyet analiz raporu"])
    if any(word in title for word in ["personel", "öğretim elemanı", "akademik"]):
        suggestions.extend(["Akademik personel özgeçmişleri", "Ders görevlendirme çizelgesi", "Eğitim-öğretim iş yükü kanıtları"])
    if any(word in title for word in ["altyapı", "laboratuvar", "tesis", "donanım"]):
        suggestions.extend(["Laboratuvar envanteri", "Bakım/kalibrasyon kayıtları", "Fotoğraflı fiziki imkan kanıtları"])
    if any(word in title for word in ["kalite", "pukö", "sürekli iyileştirme", "iyileştirme"]):
        suggestions.extend(["PUKÖ çevrimi karar ve uygulama kayıtları", "İyileştirme faaliyet planı", "Önlem/sonuç izleme tablosu"])
    suggestions.extend(["İlgili kurul kararı", "Uygulama ekran görüntüsü veya belge", "Paydaş görüşü / anket çıktısı"])
    seen: set[str] = set()
    output = []
    for item in suggestions:
        if item not in seen:
            seen.add(item)
            output.append(item)
        if len(output) >= 6:
            break
    return output




def _coach_actions(section: dict[str, Any], quality: dict[str, Any]) -> list[dict[str, str]]:
    actions: list[dict[str, str]] = []
    if quality.get("words", 0) < 120:
        actions.append({"priority": "Yüksek", "title": "Metni güçlendir", "detail": "Süreç, uygulama, ölçme sonucu ve iyileştirme kararını aynı akışta anlat."})
    if quality.get("evidence", 0) == 0:
        actions.append({"priority": "Kritik", "title": "Doğrudan kanıt bağla", "detail": "Kurul kararı, analiz raporu veya uygulama çıktısı gibi doğrulanabilir en az bir kanıt ekle."})
    if quality.get("puko", 0) < 4:
        actions.append({"priority": "Yüksek", "title": "PUKÖ döngüsünü kapat", "detail": "Planla-Uygula-Kontrol Et-Önlem Al alanlarının tamamında sonuç ve sorumlu bilgisi kullan."})
    if quality.get("tables", 0) == 0:
        actions.append({"priority": "Orta", "title": "Özet tablo ekle", "detail": "Ölçütü destekleyen sayısal veya izleme tablosu kalite skorunu yükseltir."})
    if not actions:
        actions.append({"priority": "İyileştirme", "title": "Son cümleyi güçlendir", "detail": "Metni kanıt bağlantısı ve sürdürülebilir iyileştirme sonucu ile bitir."})
    return actions[:4]


def _template_candidates(section: dict[str, Any], evidence_suggestions: list[str]) -> list[dict[str, str]]:
    title = str(section.get("section_title") or "bu ölçüt")
    main_title = str(section.get("main_title") or "rapor başlığı")
    return [
        {
            "title": "Kanıt odaklı ölçüt paragrafı",
            "kind": "Rapor Metni",
            "content": f"{title} kapsamında yürütülen faaliyetler, ilgili kurul kararları ve uygulama çıktıları ile izlenmektedir. Süreç sonuçları düzenli olarak değerlendirilmekte, elde edilen bulgular {main_title} hedefleriyle ilişkilendirilerek iyileştirme kararlarına dönüştürülmektedir.",
        },
        {
            "title": "PUKÖ kapanış cümlesi",
            "kind": "PUKÖ",
            "content": "Kontrol aşamasında elde edilen bulgular doğrultusunda eksik kalan alanlar için düzeltici/önleyici faaliyetler tanımlanmış ve bir sonraki uygulama döneminin planına aktarılmıştır.",
        },
        {
            "title": "Kanıt listesi cümlesi",
            "kind": "Kanıt",
            "content": f"Bu ölçüt için beklenen temel kanıtlar: {', '.join(evidence_suggestions[:3])}.",
        },
    ]

def _ai_suggestions(section: dict[str, Any], quality: dict[str, Any], evidence_suggestions: list[str]) -> dict[str, Any]:
    saved = _safe_json_loads(section.get("ai_suggestions_json"), {})
    if isinstance(saved, dict) and saved.get("generated_at"):
        saved.setdefault("coach_actions", _coach_actions(section, quality))
        saved.setdefault("template_suggestions", _template_candidates(section, evidence_suggestions))
        saved.setdefault("quality_explanation", _quality_explanation(quality))
        return saved
    missing = []
    if quality.get("words", 0) < 120:
        missing.append("Rapor metni kısa; süreç, uygulama, ölçme ve iyileştirme ayrıntıları genişletilmeli.")
    if quality.get("evidence", 0) == 0:
        missing.append("Bu ölçüt için en az bir doğrudan kanıt bağlanmalı.")
    if quality.get("puko", 0) < 4:
        missing.append("PUKÖ alanlarında eksik kalan başlıklar tamamlanmalı.")
    if quality.get("tables", 0) == 0:
        missing.append("Uygunsa ölçütü destekleyen kısa bir özet tablo eklenmeli.")
    if not missing:
        missing.append("Başlık temel olarak güçlü; kanıt atıfları ve sonuç cümleleri güçlendirilebilir.")
    return {
        "generated_at": "",
        "summary": "Bu başlık için hızlı kalite, kanıt ve koçluk önerileri hazır.",
        "coach_headline": "Bu bölümde kaliteyi artıracak en hızlı adımlar belirlendi.",
        "quality_explanation": _quality_explanation(quality),
        "weak_points": missing[:4],
        "coach_actions": _coach_actions(section, quality),
        "evidence_suggestions": evidence_suggestions,
        "template_suggestions": _template_candidates(section, evidence_suggestions),
        "rewrite_tips": [
            "Metinde iddia → kanıt → sonuç sırasını koru.",
            "Ölçülebilir veri, tarih ve kurul kararı referanslarını açık yaz.",
            "PUKÖ döngüsünü yalnızca faaliyet değil, sonuç/önlem ilişkisiyle anlat.",
        ],
    }


def _quality_explanation(quality: dict[str, Any]) -> list[str]:
    rows = []
    rows.append(f"Metin hacmi: {int(quality.get('words', 0) or 0)} kelime.")
    rows.append(f"Kanıt bağlantısı: {int(quality.get('evidence', 0) or 0)} kayıt.")
    rows.append(f"PUKÖ doluluğu: {int(quality.get('puko', 0) or 0)}/4 alan.")
    rows.append(f"Tablo desteği: {int(quality.get('tables', 0) or 0)} tablo.")
    return rows

def _completion_percent(section: dict[str, Any], quality: dict[str, Any]) -> int:
    score = 0
    score += min(int(quality.get("words", 0) or 0), 300) / 300 * 35
    score += min(int(quality.get("evidence", 0) or 0), 2) / 2 * 20
    score += min(int(quality.get("tables", 0) or 0), 1) * 10
    score += min(int(quality.get("puko", 0) or 0), 4) / 4 * 20
    approval = str(section.get("approval_status") or "Taslak")
    if approval == APPROVED:
        score += 15
    elif approval == SUBMITTED:
        score += 10
    elif approval == REVISION:
        score += 4
    elif section.get("status") in {READY, COMPLETED}:
        score += 7
    return max(0, min(100, round(score)))


def _risk_level(section: dict[str, Any], quality: dict[str, Any], completion: int) -> str:
    days = _days_until(section.get("deadline"))
    if section.get("approval_status") == APPROVED:
        return "good"
    if days is not None and days < 0:
        return "critical"
    if section.get("approval_status") == REVISION or section.get("status") == REVISION:
        return "critical"
    if completion < 45 or quality.get("evidence", 0) == 0:
        return "warning"
    if days is not None and 0 <= days <= 5:
        return "warning"
    return "good"


def _pro_readiness_for_card(card: dict[str, Any]) -> dict[str, Any]:
    quality_score = int(float(card.get("quality_score") or 0))
    completion = int(float(card.get("completion_percent") or 0))
    evidence_count = int(card.get("evidence_count") or 0)
    table_count = int(card.get("table_count") or 0)
    puko_done = int(card.get("puko_done") or 0)
    word_count = int(card.get("word_count") or 0)
    approval_status = str(card.get("approval_status") or "")
    checks = [
        {
            "key": "quality_98",
            "label": "Kalite skoru 98+",
            "done": quality_score >= PRO_QUALITY_TARGET,
            "weight": 22,
            "action": "Kalite skoru düşük boyutları tamamla.",
        },
        {
            "key": "completion_98",
            "label": "Tamamlanma 98+",
            "done": completion >= PRO_QUALITY_TARGET,
            "weight": 16,
            "action": "Durum, metin, kanıt, PUKÖ ve onay açıklarını kapat.",
        },
        {
            "key": "evidence_coverage",
            "label": "En az iki doğrudan kanıt",
            "done": evidence_count >= 2,
            "weight": 16,
            "action": "Bölüme ikinci doğrulanabilir kanıtı bağla.",
        },
        {
            "key": "puko_closed",
            "label": "PUKÖ döngüsü tam",
            "done": puko_done >= 4,
            "weight": 14,
            "action": "Planla, Uygula, Kontrol Et ve Önlem Al alanlarını tamamla.",
        },
        {
            "key": "deep_text",
            "label": "420+ kelime derinlik",
            "done": word_count >= 420,
            "weight": 12,
            "action": "Metne süreç, ölçüm sonucu ve iyileştirme kararını ekle.",
        },
        {
            "key": "table_support",
            "label": "Tablo/izleme verisi var",
            "done": table_count >= 1,
            "weight": 10,
            "action": "Ölçütü destekleyen tablo veya izleme verisi bağla.",
        },
        {
            "key": "approval_gate",
            "label": "Onayda veya onaylı",
            "done": approval_status in {SUBMITTED, APPROVED},
            "weight": 10,
            "action": "Başlığı onay akışına gönder veya onay kararını tamamla.",
        },
    ]
    total_weight = sum(int(item["weight"]) for item in checks)
    earned = sum(int(item["weight"]) for item in checks if item["done"])
    pro_score = round(earned / max(1, total_weight) * 100)
    blockers = [item for item in checks if not item["done"]]
    status = "ready" if not blockers else "near" if pro_score >= 82 else "work"
    label = "9.8+ hazır" if status == "ready" else "9.8+ yakında" if status == "near" else "9.8+ aksiyon gerekli"
    return {
        "target_score": PRO_QUALITY_TARGET,
        "target_label": PRO_TARGET_LABEL,
        "score": pro_score,
        "status": status,
        "label": label,
        "score_gap": max(0, PRO_QUALITY_TARGET - min(quality_score, completion)),
        "blocker_count": len(blockers),
        "checklist": checks,
        "missing_actions": [str(item["action"]) for item in blockers[:5]],
    }


def _studio_card(username: str, program_id: str, section: dict[str, Any]) -> dict[str, Any]:
    section_key = str(section.get("section_key", "") or "")
    quality = quality_for_section(username, program_id, section)
    evidence = list_evidence(username, program_id, section_key)
    tables = list_tables(username, program_id, section_key)
    evidence_suggestions = _suggested_evidence(section, quality)
    completion = _completion_percent(section, quality)
    risk = str(section.get("risk_level") or "").strip() or _risk_level(section, quality, completion)
    owner = _assigned_owner(program_id, section_key)
    ai = _ai_suggestions(section, quality, evidence_suggestions)
    card = {
        "section_key": section_key,
        "section_title": section.get("section_title", ""),
        "main_title": section.get("main_title", ""),
        "report_group_title": section.get("report_group_title", "") or section.get("main_title", ""),
        "report_subgroup_title": section.get("report_subgroup_title", ""),
        "status": section.get("status", "Başlamadı"),
        "approval_status": section.get("approval_status", "Taslak"),
        "completion_percent": completion,
        "quality_score": int(section.get("quality_score") or quality.get("score") or 0),
        "quality_dimensions": {
            "metin": min(100, round(min(int(quality.get("words", 0) or 0), 300) / 300 * 100)),
            "kanıt": min(100, round(min(int(quality.get("evidence", 0) or 0), 3) / 3 * 100)),
            "puko": min(100, round(min(int(quality.get("puko", 0) or 0), 4) / 4 * 100)),
            "tablo": min(100, 100 if int(quality.get("tables", 0) or 0) > 0 else 0),
        },
        "risk_level": risk,
        "risk_label": {"good": "İyi", "warning": "Uyarı", "critical": "Riskli"}.get(risk, "Uyarı"),
        "deadline": section.get("deadline", ""),
        "days_until_deadline": _days_until(section.get("deadline")),
        "responsible": owner,
        "last_updated_at": section.get("updated_at", ""),
        "evidence_count": len(evidence),
        "table_count": len(tables),
        "puko_done": int(quality.get("puko", 0) or 0),
        "word_count": int(quality.get("words", 0) or 0),
        "ai_ready": bool(ai.get("weak_points") or ai.get("evidence_suggestions")),
        "ai_suggestions": ai,
        "suggested_evidence": evidence_suggestions,
        "estimated_minutes": max(20, 110 - completion),
        "preview": str(section.get("report_text", "") or "")[:220],
    }
    card["pro_readiness"] = _pro_readiness_for_card(card)
    return card


def _group_tree(cards: list[dict[str, Any]]) -> list[dict[str, Any]]:
    groups: dict[str, dict[str, Any]] = {}
    for card in cards:
        group_name = str(card.get("report_group_title") or card.get("main_title") or "Diğer")
        group = groups.setdefault(group_name, {"title": group_name, "total": 0, "completed": 0, "progress_total": 0, "risk": "good", "children": {}})
        group["total"] += 1
        group["progress_total"] += int(card.get("completion_percent") or 0)
        if card.get("approval_status") == APPROVED or card.get("status") == COMPLETED:
            group["completed"] += 1
        if card.get("risk_level") == "critical":
            group["risk"] = "critical"
        elif card.get("risk_level") == "warning" and group.get("risk") != "critical":
            group["risk"] = "warning"
        subgroup_name = str(card.get("report_subgroup_title") or card.get("main_title") or "Başlıklar")
        child = group["children"].setdefault(subgroup_name, {"title": subgroup_name, "total": 0, "completed": 0, "progress_total": 0, "first_section_key": card.get("section_key"), "risk": "good"})
        child["total"] += 1
        child["progress_total"] += int(card.get("completion_percent") or 0)
        if card.get("approval_status") == APPROVED or card.get("status") == COMPLETED:
            child["completed"] += 1
        if card.get("risk_level") == "critical":
            child["risk"] = "critical"
        elif card.get("risk_level") == "warning" and child.get("risk") != "critical":
            child["risk"] = "warning"
    output = []
    for group in groups.values():
        children = []
        for child in group.pop("children", {}).values():
            child["progress"] = round(child["progress_total"] / child["total"], 1) if child["total"] else 0
            children.append(child)
        group["progress"] = round(group["progress_total"] / group["total"], 1) if group["total"] else 0
        group["children"] = children
        output.append(group)
    return output




def _studio_heatmap(cards: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for idx, card in enumerate(cards):
        score = int(card.get("quality_score") or 0)
        completion = int(card.get("completion_percent") or 0)
        risk = str(card.get("risk_level") or "warning")
        rows.append({
            "section_key": card.get("section_key"),
            "label": card.get("section_title"),
            "group": card.get("report_group_title") or card.get("main_title") or "Rapor",
            "subgroup": card.get("report_subgroup_title") or "Başlıklar",
            "score": score,
            "completion": completion,
            "risk": risk,
            "risk_label": card.get("risk_label"),
            "status": card.get("status"),
            "approval_status": card.get("approval_status"),
            "x": idx % 12,
            "y": idx // 12,
            "intensity": max(0, min(100, 100 - min(score, completion))),
        })
    return rows


def _studio_pro_overview(cards: list[dict[str, Any]]) -> dict[str, Any]:
    readiness_rows = [dict(card.get("pro_readiness") or {}) for card in cards]
    total = len(readiness_rows)
    ready = sum(1 for row in readiness_rows if row.get("status") == "ready")
    near = sum(1 for row in readiness_rows if row.get("status") == "near")
    avg_score = round(sum(int(row.get("score") or 0) for row in readiness_rows) / max(1, total), 1) if total else 0
    blockers = sum(int(row.get("blocker_count") or 0) for row in readiness_rows)
    action_rows: list[dict[str, Any]] = []
    for card in cards:
        readiness = dict(card.get("pro_readiness") or {})
        if readiness.get("status") == "ready":
            continue
        action_rows.append({
            "section_key": card.get("section_key"),
            "section_title": card.get("section_title"),
            "quality_score": card.get("quality_score", 0),
            "pro_score": readiness.get("score", 0),
            "blocker_count": readiness.get("blocker_count", 0),
            "next_action": (readiness.get("missing_actions") or ["Kalite kontrolünü tamamla."])[0],
        })
    action_rows.sort(key=lambda row: (int(row.get("pro_score") or 0), int(row.get("quality_score") or 0)))
    return {
        "target_score": PRO_QUALITY_TARGET,
        "target_label": PRO_TARGET_LABEL,
        "score": avg_score,
        "ready_sections": ready,
        "near_sections": near,
        "total_sections": total,
        "blocker_count": blockers,
        "ready_percent": round(ready / max(1, total) * 100),
        "status": "ready" if total and ready == total else "near" if avg_score >= 82 else "work",
        "summary": f"{ready}/{total} başlık 9.8+ kapısından geçti.",
        "next_actions": action_rows[:6],
    }


def _program_template_bank(username: str, program_id: str, sample_cards: list[dict[str, Any]] | None = None) -> list[dict[str, Any]]:
    tenant_id = program_tenant_id(program_id)
    rows: list[dict[str, Any]] = []
    with get_conn() as conn:
        try:
            db_rows = conn.execute(
                """SELECT id, title, content, tags, source, section_key, profile, created_by, created_at
                   FROM section_template_bank
                   WHERE COALESCE(tenant_id,'tenant_default') IN ('system', ?) AND (COALESCE(program_id,'') IN ('', ?))
                   ORDER BY CASE WHEN COALESCE(section_key,'')='' THEN 1 ELSE 0 END, created_at DESC
                   LIMIT 24""",
                (tenant_id, program_id),
            ).fetchall()
            rows = rows_to_dicts(db_rows)
        except Exception:
            rows = []
    if rows:
        return rows
    example = (sample_cards or [{}])[0] if sample_cards else {}
    section_title = str(example.get("section_title") or "ölçüt")
    return [
        {
            "id": "system-evidence-paragraph",
            "title": "Kanıtla güçlendirilmiş açıklama",
            "content": f"{section_title} kapsamında yürütülen uygulamalar ilgili kurul kararları, izleme raporları ve paydaş geri bildirimleriyle desteklenmektedir. Bulgular düzenli olarak değerlendirilmekte ve iyileştirme kararlarına dönüştürülmektedir.",
            "tags": "rapor,kanıt,puko",
            "source": "system",
        },
        {
            "id": "system-puko-closure",
            "title": "PUKÖ kapanış paragrafı",
            "content": "Kontrol aşamasında elde edilen sonuçlar doğrultusunda eksik alanlar için düzeltici/önleyici faaliyetler tanımlanmış ve sonraki dönem planına aktarılmıştır.",
            "tags": "puko,iyileştirme",
            "source": "system",
        },
        {
            "id": "system-evidence-list",
            "title": "Kanıt beklentisi ifadesi",
            "content": "Bu ölçüt için kurul kararı, uygulama çıktısı, analiz tablosu ve paydaş geri bildirim raporu gibi doğrulanabilir kanıtlar beklenmektedir.",
            "tags": "kanıt,denetim",
            "source": "system",
        },
    ]

def report_studio_payload(username: str, program_id: str) -> dict[str, Any]:
    assert_program_operation_permission(username, program_id, "report_studio.view")
    sections = list_sections(username, program_id)
    cards = [_studio_card(username, program_id, section) for section in sections]
    total = len(cards)
    ready = sum(1 for card in cards if card.get("completion_percent", 0) >= 75)
    overdue = sum(1 for card in cards if (card.get("days_until_deadline") is not None and card.get("days_until_deadline") < 0 and card.get("approval_status") != APPROVED))
    approval_waiting = sum(1 for card in cards if card.get("approval_status") == SUBMITTED)
    evidence_missing = sum(1 for card in cards if int(card.get("evidence_count") or 0) == 0)
    ai_ready = sum(1 for card in cards if card.get("ai_ready"))
    avg_completion = round(sum(int(card.get("completion_percent") or 0) for card in cards) / total, 1) if total else 0
    avg_quality = round(sum(int(card.get("quality_score") or 0) for card in cards) / total, 1) if total else 0
    critical = sum(1 for card in cards if card.get("risk_level") == "critical")
    warning = sum(1 for card in cards if card.get("risk_level") == "warning")
    pro_overview = _studio_pro_overview(cards)
    return {
        "overview": {
            "total_sections": total,
            "ready_sections": ready,
            "completion_percent": avg_completion,
            "quality_score": avg_quality,
            "critical_sections": critical,
            "warning_sections": warning,
            "overdue_sections": overdue,
            "approval_waiting": approval_waiting,
            "evidence_missing": evidence_missing,
            "ai_ready": ai_ready,
            "pro_quality_target": PRO_QUALITY_TARGET,
            "pro_quality_label": PRO_TARGET_LABEL,
            "pro_score": pro_overview["score"],
            "pro_ready_sections": pro_overview["ready_sections"],
            "estimated_finish_label": "Bu hafta" if overdue == 0 and warning < max(1, total // 4) else "Riskli / takip gerekli",
        },
        "pro_overview": pro_overview,
        "filters": {
            "mine": 0,  # Client computes exact current-user subset when it has user identity.
            "overdue": overdue,
            "approval_waiting": approval_waiting,
            "evidence_missing": evidence_missing,
            "ai_ready": ai_ready,
            "revision": sum(1 for card in cards if card.get("approval_status") == REVISION or card.get("status") == REVISION),
        },
        "cards": cards,
        "tree": _group_tree(cards),
        "heatmap": _studio_heatmap(cards),
        "template_bank": _program_template_bank(username, program_id, cards[:8]),
    }



def _program_profile(program_id: str) -> str:
    program = get_program(program_id) or {}
    return str(program.get("accreditation_profile") or program.get("profile") or "MEDEK")


def _section_full_text(section: dict[str, Any]) -> str:
    parts = [
        section.get("section_key", ""), section.get("section_title", ""), section.get("main_title", ""),
        section.get("report_text", ""), section.get("planla", ""), section.get("uygula", ""),
        section.get("kontrol", ""), section.get("onlem", ""), section.get("notes", ""),
    ]
    return " ".join(str(part or "") for part in parts).lower()


def _contains_any(text: str, keywords: list[str]) -> bool:
    value = text.lower()
    return any(keyword.lower() in value for keyword in keywords)


def _standard_expectations(section: dict[str, Any], profile: str) -> list[dict[str, Any]]:
    title_text = _section_full_text({**section, "report_text": ""})
    expectations: list[dict[str, Any]] = [
        {"id": "process", "label": "Süreç açıklığı", "keywords": ["süreç", "uygulama", "yürüt", "değerlendir"], "expected_evidence": ["Süreç akış dokümanı", "Komisyon/kurul kararı"], "description": "Bölümde uygulanan sürecin kim, ne zaman, nasıl yürüttüğü açık olmalı."},
        {"id": "evidence", "label": "Doğrudan kanıt", "keywords": ["kanıt", "ek", "belge", "tutanak", "rapor"], "expected_evidence": ["Kurul kararı", "Analiz raporu", "Uygulama çıktısı"], "description": "Metindeki iddiaları destekleyen doğrulanabilir kanıtlar bağlı olmalı."},
        {"id": "puko", "label": "PUKÖ bütünlüğü", "keywords": ["planla", "uygula", "kontrol", "önlem", "iyileştirme"], "expected_evidence": ["PUKÖ izleme tablosu", "İyileştirme karar formu"], "description": "Planla-Uygula-Kontrol Et-Önlem Al döngüsü sonuç ilişkisiyle tamamlanmalı."},
        {"id": "measurement", "label": "Ölçülebilir sonuç", "keywords": ["gösterge", "ölç", "oran", "sonuç", "analiz", "%"], "expected_evidence": ["Performans göstergesi tablosu", "Anket/ölçme sonuç raporu"], "description": "Akreditasyon metni nicel/nitel sonuç ve gösterge ile desteklenmeli."},
    ]
    if _contains_any(title_text, ["program amaç", "hedef", "misyon", "vizyon"]):
        expectations.append({"id": "stakeholder", "label": "Paydaş katkısı", "keywords": ["paydaş", "mezun", "işveren", "danışma", "görüş"], "expected_evidence": ["Danışma kurulu tutanağı", "Paydaş görüş formu"], "description": "Amaç/hedef güncellemelerinde iç ve dış paydaş katkısı gösterilmeli."})
    if _contains_any(title_text, ["ders", "müfredat", "program çıktı", "öğrenme çıktı"]):
        expectations.append({"id": "curriculum", "label": "Müfredat ve çıktı ilişkisi", "keywords": ["program çıkt", "ders", "matris", "müfredat", "öğrenme"], "expected_evidence": ["Ders-program çıktısı matrisi", "Ders bilgi paketi"], "description": "Dersler, program çıktıları ve ölçme-değerlendirme ilişkisi görünür olmalı."})
    if _contains_any(title_text, ["öğrenci", "mezun", "danışman"]):
        expectations.append({"id": "student_feedback", "label": "Öğrenci/mezun geri bildirimi", "keywords": ["öğrenci", "mezun", "anket", "geri bildirim", "memnuniyet"], "expected_evidence": ["Öğrenci memnuniyet anketi", "Mezun izleme raporu"], "description": "Öğrenci veya mezun verisiyle izleme ve iyileştirme ilişkisi kurulmalı."})
    if _contains_any(title_text, ["laboratuvar", "altyapı", "tesis", "donanım"]):
        expectations.append({"id": "infrastructure", "label": "Altyapı yeterliliği", "keywords": ["laboratuvar", "donanım", "envanter", "bakım", "kalibrasyon"], "expected_evidence": ["Laboratuvar envanteri", "Bakım/kalibrasyon kaydı", "Fotoğraflı kanıt"], "description": "Fiziki/teknik altyapı güncel envanter ve kullanım kanıtlarıyla desteklenmeli."})
    if profile.upper() in {"MEDEK", "MÜDEK", "MUDEK"}:
        expectations.append({"id": "continuous_improvement", "label": "Sürekli iyileştirme izi", "keywords": ["iyileştirme", "önlem", "aksiyon", "sonraki dönem", "izleme"], "expected_evidence": ["İyileştirme aksiyon planı", "Önlem takip tablosu"], "description": "Bölüm sadece mevcut durumu değil, iyileştirme kararını ve izleme sonucunu da göstermeli."})
    return expectations


def accreditation_gap_scan(username: str, program_id: str, section_key: str) -> dict[str, Any]:
    assert_program_operation_permission(username, program_id, "report_studio.standards_scan")
    section = get_section(username, program_id, section_key)
    if not section:
        raise KeyError(section_key)
    profile = _program_profile(program_id)
    quality = quality_for_section(username, program_id, section)
    evidence = list_evidence(username, program_id, section_key)
    tables = list_tables(username, program_id, section_key)
    full_text = _section_full_text(section)
    expectations = _standard_expectations(section, profile)
    findings: list[dict[str, Any]] = []
    for exp in expectations:
        text_ok = _contains_any(full_text, exp["keywords"])
        evidence_ok = bool(evidence) and any(_contains_any(f"{item.get('code','')} {item.get('original_name','')} {item.get('note','')}", exp["expected_evidence"] + exp["keywords"]) for item in evidence)
        table_ok = bool(tables) if exp["id"] in {"measurement", "puko", "continuous_improvement"} else True
        status = "tamam" if text_ok and evidence_ok and table_ok else "eksik" if not text_ok and not evidence_ok else "zayıf"
        severity = "kritik" if status == "eksik" and exp["id"] in {"evidence", "puko", "process"} else "uyarı" if status != "tamam" else "iyi"
        missing = []
        if not text_ok:
            missing.append("Metinde standart beklentisini karşılayan açık ifade yok.")
        if not evidence_ok:
            missing.append("Beklenen kanıt türü bağlı değil veya kanıt notunda görünmüyor.")
        if not table_ok:
            missing.append("Bu beklenti için destekleyici tablo/izleme verisi önerilir.")
        findings.append({
            "id": exp["id"],
            "label": exp["label"],
            "description": exp["description"],
            "status": status,
            "severity": severity,
            "missing": missing,
            "expected_evidence": exp["expected_evidence"],
            "quick_fix": f"{exp['label']} için metne kısa bir uygulama-sonuç cümlesi ve en az bir doğrudan kanıt bağlantısı ekleyin.",
        })
    score = max(0, min(100, int(quality.get("score", 0) or 0) - (sum(1 for f in findings if f["severity"] == "kritik") * 8) - (sum(1 for f in findings if f["severity"] == "uyarı") * 4)))
    result = {
        "mode": "standards_gap_scan",
        "generated_at": now_iso(),
        "profile": profile,
        "section_key": section_key,
        "section_title": section.get("section_title", ""),
        "score": score,
        "summary": f"{profile} ölçüt beklentilerine göre {sum(1 for f in findings if f['status']=='tamam')}/{len(findings)} kontrol güçlü görünüyor.",
        "critical_count": sum(1 for f in findings if f["severity"] == "kritik"),
        "warning_count": sum(1 for f in findings if f["severity"] == "uyarı"),
        "findings": findings,
        "next_actions": [f["quick_fix"] for f in findings if f["status"] != "tamam"][:5],
    }
    log_activity("Standart eksiklik tarama", f"{section_key} için akreditasyon standart taraması üretildi", username, program_id)
    return result


def evidence_matching_assistant(username: str, program_id: str, section_key: str) -> dict[str, Any]:
    assert_program_operation_permission(username, program_id, "report_studio.evidence_match")
    section = get_section(username, program_id, section_key)
    if not section:
        raise KeyError(section_key)
    quality = quality_for_section(username, program_id, section)
    evidence = list_evidence(username, program_id, section_key)
    expectations = _standard_expectations(section, _program_profile(program_id))
    section_text = _section_full_text(section)
    linked: list[dict[str, Any]] = []
    for item in evidence:
        blob = f"{item.get('code','')} {item.get('original_name','')} {item.get('note','')}".lower()
        matched_expectations = [exp["label"] for exp in expectations if _contains_any(blob + " " + section_text, exp["keywords"] + exp["expected_evidence"])]
        note_ok = bool(str(item.get("note") or "").strip())
        linked.append({
            "id": item.get("id"),
            "code": item.get("code"),
            "name": item.get("original_name"),
            "note": item.get("note"),
            "match_score": min(100, 45 + len(matched_expectations) * 15 + (20 if note_ok else 0)),
            "status": "güçlü" if len(matched_expectations) >= 2 and note_ok else "zayıf" if matched_expectations else "belirsiz",
            "matched_expectations": matched_expectations[:4],
            "recommendation": "Kanıt notuna hangi bulguyu desteklediğini yazın." if not note_ok else "Metinde bu kanıta açık atıf yapın.",
        })
    expected_pool: list[str] = []
    for exp in expectations:
        expected_pool.extend(exp["expected_evidence"])
    existing_blob = " ".join(f"{item.get('original_name','')} {item.get('code','')} {item.get('note','')}" for item in evidence).lower()
    missing = []
    seen = set()
    for item in expected_pool:
        if item not in seen and item.lower() not in existing_blob:
            seen.add(item)
            missing.append(item)
        if len(missing) >= 8:
            break
    result = {
        "mode": "evidence_matching",
        "generated_at": now_iso(),
        "section_key": section_key,
        "linked_count": len(linked),
        "strong_count": sum(1 for item in linked if item["status"] == "güçlü"),
        "weak_count": sum(1 for item in linked if item["status"] != "güçlü"),
        "coverage_score": min(100, int(quality.get("evidence", 0) or 0) * 25 + sum(1 for item in linked if item["status"] == "güçlü") * 15),
        "summary": "Bağlı kanıtlar standart beklentileriyle eşleştirildi; zayıf kanıtlar ve önerilen yeni kanıtlar listelendi.",
        "linked_evidence": linked,
        "missing_evidence": missing,
        "recommended_links": [
            {"expected": item, "reason": "Bu kanıt türü seçili ölçütte denetim sırasında beklenebilir."}
            for item in missing[:5]
        ],
    }
    log_activity("Kanıt eşleştirme asistanı", f"{section_key} için kanıt eşleştirme analizi üretildi", username, program_id)
    return result

def quick_ai_suggestions(username: str, program_id: str, section_key: str, mode: str = "coach") -> dict[str, Any]:
    assert_program_operation_permission(username, program_id, "ai.local.draft")
    section = get_section(username, program_id, section_key)
    if not section:
        raise KeyError(section_key)
    quality = quality_for_section(username, program_id, section)
    evidence_suggestions = _suggested_evidence(section, quality)
    suggestions = _ai_suggestions(section, quality, evidence_suggestions)
    suggestions["generated_at"] = now_iso()
    suggestions["mode"] = mode
    if mode == "puko":
        title = section.get("section_title", "bu ölçüt")
        suggestions["puko"] = {
            "planla": f"{title} kapsamında mevcut durum, paydaş beklentileri ve hedef göstergeler analiz edilerek yıllık iyileştirme planı oluşturulur.",
            "uygula": "Planlanan faaliyetler sorumlu kişiler, terminler ve kanıt kayıtları ile uygulanır; ilgili kurul/komisyon kararları dosyalanır.",
            "kontrol": "Uygulama sonuçları performans göstergeleri, anketler, toplantı tutanakları ve kanıt yeterliliği üzerinden dönem sonunda değerlendirilir.",
            "onlem": "Eksik veya zayıf kalan alanlar için düzeltici/önleyici faaliyetler tanımlanır ve bir sonraki dönemin planına aktarılır.",
        }
    with transaction() as conn:
        conn.execute(
            "UPDATE sections SET ai_suggestions_json=?, last_ai_review_at=?, updated_at=? WHERE program_id=? AND section_key=?",
            (json.dumps(suggestions, ensure_ascii=False), suggestions["generated_at"], now_iso(), program_id, section_key),
        )
    log_activity("AI hızlı öneri", f"{section_key} için {mode} önerisi üretildi", username, program_id)
    return suggestions


def recalculate_section_quality(username: str, program_id: str, section_key: str) -> dict[str, Any]:
    section = get_section(username, program_id, section_key)
    if not section:
        raise KeyError(section_key)
    quality = quality_for_section(username, program_id, section)
    completion = _completion_percent(section, quality)
    risk = _risk_level(section, quality, completion)
    with transaction() as conn:
        conn.execute(
            "UPDATE sections SET quality_score=?, risk_level=?, updated_at=? WHERE program_id=? AND section_key=?",
            (int(quality.get("score", 0) or 0), risk, now_iso(), program_id, section_key),
        )
    return {"section_key": section_key, "quality_score": int(quality.get("score", 0) or 0), "completion_percent": completion, "risk_level": risk}


def bulk_studio_update(username: str, program_id: str, payload: dict[str, Any]) -> dict[str, Any]:
    keys = [str(key).strip() for key in payload.get("section_keys", []) if str(key).strip()]
    if not keys:
        raise ValueError("En az bir başlık seçin.")
    action = str(payload.get("action", "") or "").strip()
    status = str(payload.get("status", "") or "").strip()
    deadline = str(payload.get("deadline", "") or "").strip()
    append_note = str(payload.get("note", "") or "").strip()
    updated = 0
    skipped: list[str] = []
    if action in {"status", "deadline", "status_deadline"}:
        permission = "bulk.status.update" if status else "bulk.deadline.update"
        assert_program_operation_permission(username, program_id, permission)
    elif action in {"ai_draft", "puko_draft", "quality"}:
        assert_program_operation_permission(username, program_id, "report_studio.bulk_ai" if action in {"ai_draft", "puko_draft"} else "full_report.quality.view")
    else:
        raise ValueError("Desteklenmeyen toplu işlem.")

    if status and status not in STATUS_OPTIONS:
        raise ValueError("Geçersiz durum seçimi.")

    for key in keys:
        try:
            section = get_section(username, program_id, key)
            if not section:
                skipped.append(key)
                continue
            if action == "ai_draft":
                suggestions = quick_ai_suggestions(username, program_id, key, "coach")
                note = append_note or "AI taslağı/önerisi hazırlandı. Sağ panelden incelenebilir."
                current_notes = str(section.get("notes", "") or "")
                update_section(username, program_id, key, {**section, "notes": f"{current_notes}\n\n{note}".strip()})
            elif action == "quality":
                recalculate_section_quality(username, program_id, key)
            else:
                patch = dict(section)
                if status:
                    patch["status"] = status
                if deadline:
                    patch["deadline"] = deadline
                if append_note:
                    patch["notes"] = f"{str(section.get('notes','') or '').strip()}\n\n{append_note}".strip()
                update_section(username, program_id, key, patch)
            updated += 1
        except PermissionError:
            skipped.append(key)
        except KeyError:
            skipped.append(key)
    log_activity("Rapor Stüdyosu toplu işlem", f"{action}: {updated}/{len(keys)}", username, program_id)
    return {"requested": len(keys), "updated": updated, "skipped": skipped, "action": action}




def section_template_bank(username: str, program_id: str, section_key: str = "") -> list[dict[str, Any]]:
    assert_program_operation_permission(username, program_id, "report_studio.template_bank.view")
    cards = []
    if section_key:
        section = get_section(username, program_id, section_key)
        if not section:
            raise KeyError(section_key)
        cards = [_studio_card(username, program_id, section)]
    return _program_template_bank(username, program_id, cards)


def create_section_template(username: str, program_id: str, payload: dict[str, Any]) -> dict[str, Any]:
    assert_program_operation_permission(username, program_id, "report_studio.template_bank.manage")
    title = str(payload.get("title") or "").strip()
    content = str(payload.get("content") or "").strip()
    if not title or not content:
        raise ValueError("Şablon başlığı ve içeriği zorunludur.")
    section_key = str(payload.get("section_key") or "").strip()
    if section_key and not get_section(username, program_id, section_key):
        raise KeyError(section_key)
    item = {
        "id": str(uuid.uuid4()),
        "tenant_id": program_tenant_id(program_id),
        "program_id": program_id,
        "section_key": section_key,
        "profile": str(payload.get("profile") or ""),
        "title": title,
        "content": content,
        "tags": str(payload.get("tags") or ""),
        "source": "manual",
        "created_by": username,
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    with transaction() as conn:
        conn.execute(
            """INSERT INTO section_template_bank(id, tenant_id, program_id, section_key, profile, title, content, tags, source, created_by, created_at, updated_at)
               VALUES(?,?,?,?,?,?,?,?,?,?,?,?)""",
            (item["id"], item["tenant_id"], item["program_id"], item["section_key"], item["profile"], item["title"], item["content"], item["tags"], item["source"], item["created_by"], item["created_at"], item["updated_at"]),
        )
    log_activity("Şablon bankası", f"{title} şablonu eklendi", username, program_id)
    return item

def collaboration_ping(username: str, program_id: str, section_key: str) -> list[dict[str, Any]]:
    assert_program_operation_permission(username, program_id, "report_studio.collaboration.view")
    section = get_section(username, program_id, section_key)
    if not section:
        raise KeyError(section_key)
    user = get_user(username, active_only=True) or {}
    display_name = str(user.get("full_name") or username)
    ts = now_iso()
    with transaction() as conn:
        conn.execute(
            """INSERT INTO section_collaboration_sessions(id, program_id, section_key, username, display_name, started_at, last_seen_at, status, tenant_id)
               VALUES(?,?,?,?,?,?,?,?,?)
               ON CONFLICT(program_id, section_key, username) DO UPDATE SET display_name=excluded.display_name, last_seen_at=excluded.last_seen_at, status='editing'""",
            (str(uuid.uuid4()), program_id, section_key, username, display_name, ts, ts, "editing", program_tenant_id(program_id)),
        )
    return active_collaborators(username, program_id, section_key)


def active_collaborators(username: str, program_id: str, section_key: str) -> list[dict[str, Any]]:
    assert_program_operation_permission(username, program_id, "report_studio.collaboration.view")
    section = get_section(username, program_id, section_key)
    if not section:
        raise KeyError(section_key)
    cutoff = (datetime.now() - timedelta(minutes=2)).strftime("%Y-%m-%d %H:%M:%S")
    with get_conn() as conn:
        rows = conn.execute(
            """SELECT username, display_name, last_seen_at, status
               FROM section_collaboration_sessions
               WHERE program_id=? AND section_key=? AND last_seen_at>=?
               ORDER BY last_seen_at DESC""",
            (program_id, section_key, cutoff),
        ).fetchall()
    return rows_to_dicts(rows)


def section_docx_bytes(username: str, program_id: str, section_key: str) -> bytes:
    assert_program_operation_permission(username, program_id, "report_studio.section_export")
    section = get_section(username, program_id, section_key)
    if not section:
        raise KeyError(section_key)
    evidence = list_evidence(username, program_id, section_key)
    tables = list_tables(username, program_id, section_key)
    quality = quality_for_section(username, program_id, section)
    doc = Document()
    doc.add_heading(f"{section.get('section_key', '')} - {section.get('section_title', '')}", level=1)
    doc.add_paragraph(f"Ana başlık: {section.get('main_title', '')}")
    doc.add_paragraph(f"Durum: {section.get('status', '')} / {section.get('approval_status', '')}")
    doc.add_paragraph(f"Kalite skoru: {quality.get('score', 0)}/100")
    doc.add_heading("Rapor Metni", level=2)
    doc.add_paragraph(str(section.get('report_text', '') or ''))
    doc.add_heading("PUKÖ", level=2)
    labels = [('planla', 'Planla'), ('uygula', 'Uygula'), ('kontrol', 'Kontrol Et'), ('onlem', 'Önlem Al')]
    for key, label in labels:
        doc.add_paragraph(label, style='List Bullet')
        doc.add_paragraph(str(section.get(key, '') or ''))
    doc.add_heading("Kanıtlar", level=2)
    if evidence:
        for item in evidence:
            doc.add_paragraph(f"{item.get('code','')} - {item.get('original_name','')} ({item.get('note','')})", style='List Bullet')
    else:
        doc.add_paragraph("Bağlı kanıt yok.")
    doc.add_heading("Tablolar", level=2)
    if tables:
        for item in tables:
            doc.add_paragraph(str(item.get('table_name','')), style='List Bullet')
    else:
        doc.add_paragraph("Bağlı tablo yok.")
    output = BytesIO()
    doc.save(output)
    return output.getvalue()
