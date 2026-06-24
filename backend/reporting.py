from __future__ import annotations

import os
import re
import shutil
import subprocess
import sys
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .file_security import safe_stored_path, slugify
from .repositories import (
    control_rows,
    EPDAD_REPORT_GROUP_TITLE,
    evidence_links,
    get_settings,
    list_evidence,
    list_sections,
    list_tables,
    record_export,
    stats_payload,
)

REPORT_TEMPLATE_REVISION = "REV.2025V1.1"

REPORT_MAIN_HEADINGS = {
    "Programa İlişkin Genel Bilgiler": "A. Programa İlişkin Genel Bilgiler",
    "Ek I. Programa İlişkin Ek Bilgiler": "EK I – PROGRAMA İLİŞKİN EK BİLGİLER",
    "Ek II. Kurum Profili": "EK II – KURUM PROFİLİ",
    "Ölçüt 9. Disipline Özgü Ölçütler": "Ölçüt 9. Programa Özgü Ölçütler",
}

REPORT_REQUIRED_TABLES = {
    "1.1.2": ["Tablo 1.1. Öğrencilerin Üniversite Giriş Sınav Derecelerine İlişkin Bilgi"],
    "1.2": ["Tablo 1.2. Kayıtlı Öğrenci ve Mezun Sayıları"],
    "1.3": ["Tablo 1.3 Yatay Geçiş, Çift Anadal, Yandal Yapan Öğrenci Sayıları"],
    "5.1": ["Tablo 5.1. Öğretim Planı"],
    "5.2": ["Tablo 5.2 Ders ve Sınıf Büyüklükleri"],
    "5.3": ["Tablo 5.3. Programa/alana özgü öğrenim çıktılarını sağlayan mesleki dersler"],
    "6.1": ["Tablo 6.1. Öğretim Kadrosunun Analizi"],
    "6.4": ["Tablo 6.2. Öğretim Kadrosu Yük Özeti"],
    "EK-I.1": ["DERS BİLGİ PAKETİ"],
    "EK-II.2": [
        "Tablo II.2a Programın destek verdiği birimler",
        "Tablo II.2b Programın destek aldığı birimler",
    ],
    "EK-II.3": [
        "Tablo II.3. Personel Sayısı",
        "Tablo II-4 Öğrenci ve Mezun Sayıları",
    ],
}


CRITERIA_REPORT_GROUP_TITLES = {"B. Değerlendirme Özeti", EPDAD_REPORT_GROUP_TITLE}


def _is_criteria_group_title(group_title: str) -> bool:
    return group_title in CRITERIA_REPORT_GROUP_TITLES


def _is_criteria_summary_title(main_title: str) -> bool:
    return main_title.startswith(("Ölçüt ", "ES "))


def _apply_doc_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size in [("Heading 1", 13), ("Heading 2", 11), ("Heading 3", 10)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _apply_header_footer(doc: Document, settings: dict[str, str]) -> None:
    report_short = _clean_text(settings.get("report_short")) or "ÖDR"
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.text = ""
        run = header.add_run(
            f"Doküman No: {settings.get('report_no', 'RP-002')}\n"
            f"İlk Yayın Tarihi: {settings.get('doc_date', '')}\n"
            f"Revizyon Tarihi: {settings.get('rev_date', '')}\n"
            f"Revizyon No: {settings.get('rev_no', '01')}"
        )
        run.bold = True
        run.font.size = Pt(8)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = ""
        footer.add_run(f"{settings.get('program', '')} {report_short} | Sayfa: ").font.size = Pt(8)
        _add_page_number(footer)


def _add_hyperlink(paragraph, text: str, target: str, color: str = "0563C1") -> None:
    if not target:
        paragraph.add_run(text)
        return
    rel_id = paragraph.part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run_elm = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color_elm = OxmlElement("w:color")
    color_elm.set(qn("w:val"), color)
    rpr.append(color_elm)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run_elm.append(rpr)
    text_elm = OxmlElement("w:t")
    text_elm.text = text
    run_elm.append(text_elm)
    hyperlink.append(run_elm)
    paragraph._p.append(hyperlink)


def _evidence_target(evidence: dict[str, Any], program_id: str = "", base_url: str = "") -> str:
    evidence_id = str(evidence.get("id", "") or "").strip()
    if base_url and program_id and evidence_id:
        return f"{base_url.rstrip('/')}/api/programs/{program_id}/evidence/{evidence_id}/download"
    path = safe_stored_path(str(evidence.get("stored_path", "") or ""))
    if not path or not path.exists():
        return ""
    try:
        return path.resolve().as_uri()
    except Exception:
        return ""


def _add_paragraph_with_evidence_links(doc: Document, text: str, targets: dict[str, str]) -> None:
    pattern = re.compile(r"(?<![\w.-])((?:EK-[IVX]+(?:\.\d+)*|[A-Z](?:\.\d+)*|\d+(?:\.\d+)*)\.K\d+(?:\.\d+)?)(?![\w.-])", re.IGNORECASE)
    for chunk in _body_text_chunks(text):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(0.75)
        paragraph.paragraph_format.space_after = Pt(6)
        pos = 0
        for match in pattern.finditer(chunk):
            if match.start() > pos:
                paragraph.add_run(chunk[pos:match.start()])
            code = match.group(1)
            _add_hyperlink(paragraph, code, targets.get(code.lower(), ""))
            pos = match.end()
        if pos < len(chunk):
            paragraph.add_run(chunk[pos:])


def _format_cell(cell, font_size: int = 8, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(font_size)
            run.bold = bold


def _hex_color(value: str, fallback: str = "FFFFFF") -> str:
    clean = str(value or "").strip().lstrip("#")
    if re.fullmatch(r"[0-9a-fA-F]{6}", clean):
        return clean.upper()
    return fallback


def _set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), _hex_color(color, "FFFFFF"))


def _set_cell_borders(cell, color: str = "D7E3F1") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), _hex_color(color, "D7E3F1"))


def _align_from_value(value: str):
    return {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(str(value or "").lower(), WD_ALIGN_PARAGRAPH.LEFT)


def _apply_rich_cell_style(cell, style: dict[str, Any], options: dict[str, Any], header: bool = False) -> None:
    align = _align_from_value(str(style.get("align") or ("center" if header else options.get("align", "left"))))
    try:
        font_size = int(style.get("fontSize") or options.get("fontSize") or 8)
    except Exception:
        font_size = 8
    _format_cell(cell, font_size=font_size, bold=bool(style.get("bold") or header), align=align)
    _set_cell_borders(cell, str(options.get("borderColor") or "D7E3F1"))
    background = style.get("bg") or (options.get("headerBg") if header else "")
    if background:
        _set_cell_shading(cell, str(background))
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.italic = bool(style.get("italic"))
            run.underline = bool(style.get("underline"))
            if style.get("color"):
                run.font.color.rgb = RGBColor.from_string(_hex_color(str(style.get("color")), "142037"))


def _table_columns(rows: list[dict[str, Any]], meta: dict[str, Any]) -> list[str]:
    if isinstance(meta.get("columns"), list):
        columns = [str(col) for col in meta.get("columns", []) if str(col).strip()]
        if columns:
            return columns
    columns: list[str] = []
    for row in rows:
        for key in row.keys():
            if key not in columns:
                columns.append(str(key))
    return columns


def _span_value(style: dict[str, Any], key: str, maximum: int) -> int:
    try:
        value = int(style.get(key) or 1)
    except Exception:
        value = 1
    return max(1, min(maximum, value))


def _add_table(doc: Document, rows: list[dict[str, Any]], title: str = "", meta: dict[str, Any] | None = None) -> None:
    meta = meta if isinstance(meta, dict) else {}
    options = meta.get("options", {}) if isinstance(meta.get("options"), dict) else {}
    cell_styles = meta.get("cells", {}) if isinstance(meta.get("cells"), dict) else {}
    if title:
        paragraph = doc.add_paragraph()
        paragraph.add_run(title).bold = True
    if not rows:
        doc.add_paragraph("[Tablo verisi bulunmamaktadır.]")
        return
    columns = _table_columns(rows, meta)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.autofit = True
    for idx, column in enumerate(columns):
        table.rows[0].cells[idx].text = column
        _apply_rich_cell_style(table.rows[0].cells[idx], {}, options, header=True)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, column in enumerate(columns):
            style = cell_styles.get(f"{row_idx}:{idx}", {})
            style = style if isinstance(style, dict) else {}
            cells[idx].text = "" if style.get("hidden") else str(row.get(column, "") or "")
            _apply_rich_cell_style(cells[idx], style, options)
    for row_idx, row in enumerate(rows):
        for col_idx, column in enumerate(columns):
            style = cell_styles.get(f"{row_idx}:{col_idx}", {})
            style = style if isinstance(style, dict) else {}
            if style.get("hidden"):
                continue
            colspan = _span_value(style, "colspan", len(columns) - col_idx)
            rowspan = _span_value(style, "rowspan", len(rows) - row_idx)
            if colspan > 1 or rowspan > 1:
                try:
                    start = table.cell(row_idx + 1, col_idx)
                    end = table.cell(row_idx + rowspan, col_idx + colspan - 1)
                    start.merge(end)
                    start.text = str(row.get(column, "") or "")
                    _apply_rich_cell_style(start, style, options)
                except Exception:
                    continue
    doc.add_paragraph()


def _clean_text(value: Any) -> str:
    return str(value or "").strip()


def _fold_body_text(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", str(value or "").casefold().replace("ı", "i")).strip()


def _is_body_noise_line(line: str) -> bool:
    folded = _fold_body_text(line)
    noise_prefixes = (
        "dokuman no",
        "ilk yayin tarihi",
        "revizyon tarihi",
        "revizyon no",
        "sayfa",
        "medek mesleki egitim degerlendirme",
        "32 evler mahallesi",
    )
    return any(folded.startswith(prefix) for prefix in noise_prefixes)


def _is_body_table_artifact_line(line: str) -> bool:
    clean = _clean_text(line)
    if clean.count("|") >= 2:
        return True
    cells = [cell for cell in re.split(r"\t+| {3,}", str(line or "").strip()) if cell.strip()]
    if len(cells) >= 4 and sum(1 for cell in cells if len(cell.strip()) <= 24) >= 3:
        return True
    return False


def _starts_body_list_item(line: str) -> bool:
    return bool(re.match(r"^\s*(?:[-*•●]\s+|\d{1,2}[.)]\s+|[a-zçğıöşü][.)]\s+)", line, re.IGNORECASE))


def _body_text_chunks(text: str) -> list[str]:
    raw_lines = str(text or "").replace("\r", "\n").split("\n")
    lines: list[str] = []
    for raw_line in raw_lines:
        line = re.sub(r"\s+", " ", raw_line.replace("\u00a0", " ")).strip()
        if not line:
            if lines and lines[-1] != "":
                lines.append("")
            continue
        if _is_body_noise_line(line) or _is_body_table_artifact_line(raw_line):
            continue
        lines.append(line)

    nonblank = [line for line in lines if line]
    if not nonblank:
        return []
    wrapped_like = len(nonblank) >= 3 and (
        sum(1 for line in nonblank if 20 <= len(line) <= 130) >= max(2, int(len(nonblank) * 0.55))
        or sum(1 for line in nonblank if not re.search(r"[.!?:;)]$", line)) >= int(len(nonblank) * 0.35)
    )
    if not wrapped_like:
        return nonblank

    chunks: list[str] = []
    current: list[str] = []
    for line in lines:
        if not line:
            if current:
                chunks.append(" ".join(current).strip())
                current = []
            continue
        if current and _starts_body_list_item(line):
            chunks.append(" ".join(current).strip())
            current = [line]
        else:
            current.append(line)
    if current:
        chunks.append(" ".join(current).strip())
    return [chunk for chunk in chunks if chunk]


def _report_main_heading(main_title: str) -> str:
    return REPORT_MAIN_HEADINGS.get(_clean_text(main_title), _clean_text(main_title))


def _program_cover_label(settings: dict[str, str]) -> str:
    program = _clean_text(settings.get("program"))
    if not program:
        return ""
    if re.search(r"\bprogram[ıi]\b", program, re.IGNORECASE):
        return program
    return f"{program} Programı"


def _add_centered_paragraph(doc: Document, text: str, size: int, bold: bool = True, space_after: int = 8) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def _add_template_cover(doc: Document, settings: dict[str, str]) -> None:
    for _ in range(5):
        doc.add_paragraph()
    _add_centered_paragraph(doc, _clean_text(settings.get("report_type")) or "ÖZ DEĞERLENDİRME RAPORU", 22)
    if _clean_text(settings.get("accreditation_label")):
        _add_centered_paragraph(doc, _clean_text(settings.get("accreditation_label")), 13)
    for _ in range(2):
        doc.add_paragraph()
    for text, size in [
        (_clean_text(settings.get("university")), 16),
        (_clean_text(settings.get("school")), 15),
        (_clean_text(settings.get("department")), 13),
        (_program_cover_label(settings), 15),
        (_clean_text(settings.get("report_year")), 14),
        (_clean_text(settings.get("template_revision")) or _clean_text(settings.get("rev_code")) or REPORT_TEMPLATE_REVISION, 11),
    ]:
        if text:
            _add_centered_paragraph(doc, text, size)
    doc.add_page_break()


def _section_heading_text(section: dict[str, Any]) -> str:
    key = _clean_text(section.get("section_key"))
    title = _clean_text(section.get("section_title"))
    if not key:
        return title
    if key.startswith("A."):
        return f"{key}. {title}"
    if key.startswith("EK-I."):
        suffix = key.replace("EK-I.", "I.", 1)
        return f"{suffix} {title}"
    if key.startswith("EK-II."):
        suffix = key.replace("EK-II.", "II.", 1)
        return f"{suffix} {title}"
    return f"{key}. {title}"


def _group_report_sections(sections: list[dict[str, Any]]) -> list[tuple[str, list[dict[str, Any]]]]:
    groups: list[tuple[str, list[dict[str, Any]]]] = []
    index_by_title: dict[str, int] = {}
    for section in sections:
        group_title = _clean_text(section.get("report_group_title")) or _report_main_heading(_clean_text(section.get("main_title")))
        if group_title not in index_by_title:
            index_by_title[group_title] = len(groups)
            groups.append((group_title, []))
        groups[index_by_title[group_title]][1].append(section)
    return groups


def _group_subsections(sections: list[dict[str, Any]]) -> list[tuple[str, list[dict[str, Any]]]]:
    groups: list[tuple[str, list[dict[str, Any]]]] = []
    index_by_title: dict[str, int] = {}
    for section in sections:
        subgroup = _clean_text(section.get("report_subgroup_title")) or _clean_text(section.get("main_title"))
        if subgroup not in index_by_title:
            index_by_title[subgroup] = len(groups)
            groups.append((subgroup, []))
        groups[index_by_title[subgroup]][1].append(section)
    return groups


def _add_contents(doc: Document, sections: list[dict[str, Any]], report_type: str = "ÖZ DEĞERLENDİRME RAPORU") -> None:
    doc.add_heading("İÇİNDEKİLER", level=1)
    doc.add_paragraph("KAPAK")
    doc.add_paragraph(_clean_text(report_type) or "ÖZ DEĞERLENDİRME RAPORU")
    for group_title, group_sections in _group_report_sections(sections):
        paragraph = doc.add_paragraph()
        paragraph.add_run(group_title).bold = True
        if _is_criteria_group_title(group_title):
            for subgroup, subgroup_sections in _group_subsections(group_sections):
                child = doc.add_paragraph()
                child.paragraph_format.left_indent = Cm(0.6)
                child.add_run(subgroup).bold = True
                for section in subgroup_sections:
                    grandchild = doc.add_paragraph()
                    grandchild.paragraph_format.left_indent = Cm(1.2)
                    grandchild.add_run(_section_heading_text(section))
        else:
            for section in group_sections:
                child = doc.add_paragraph()
                child.paragraph_format.left_indent = Cm(0.6)
                child.add_run(_section_heading_text(section))
    doc.add_page_break()


def _add_empty_note(doc: Document, text: str = "[Bu başlık henüz doldurulmamıştır.]") -> None:
    run = doc.add_paragraph().add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)


def _add_puko_summary(doc: Document, section: dict[str, Any], heading_level: int = 3) -> None:
    puko_parts = [
        ("Planla", section.get("planla", "")),
        ("Uygula", section.get("uygula", "")),
        ("Kontrol Et", section.get("kontrol", "")),
        ("Önlem Al", section.get("onlem", "")),
    ]
    if not any(_clean_text(value) for _, value in puko_parts):
        return
    doc.add_heading("PUKÖ Döngüsü Özeti", level=min(heading_level, 3))
    for label, value in puko_parts:
        if _clean_text(value):
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(str(value))


def _table_title_matches(actual: str, required: str) -> bool:
    def normalize(value: str) -> str:
        return re.sub(r"\s+", " ", value.casefold().replace(".", "").replace("–", "-")).strip()

    actual_norm = normalize(actual)
    required_norm = normalize(required)
    return required_norm in actual_norm or actual_norm in required_norm


def _add_required_table_markers(doc: Document, section_key: str, tables: list[dict[str, Any]]) -> None:
    required_titles = REPORT_REQUIRED_TABLES.get(section_key, [])
    if not required_titles:
        return
    table_names = [_clean_text(table.get("table_name")) for table in tables]
    for title in required_titles:
        if any(_table_title_matches(name, title) for name in table_names if name):
            continue
        paragraph = doc.add_paragraph()
        paragraph.add_run(title).bold = True
        _add_empty_note(doc, "[Bu zorunlu şablon tablosu için Tablo Yönetimi ekranından veri eklenmelidir.]")


def _add_report_summary(doc: Document, username: str, program_id: str) -> None:
    payload = stats_payload(username, program_id)
    summary = payload.get("summary", {})
    criteria = payload.get("criteria", [])
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.add_run("Programın öz değerlendirme durumu; hazırlık, onay ve kalite göstergeleri üzerinden özetlenmiştir. ")
    paragraph.add_run(
        f"{summary.get('ready_sections', 0)}/{summary.get('total_sections', 0)} başlık hazır, "
        f"{summary.get('approved_sections', 0)} başlık onaylı, "
        f"{summary.get('submitted_sections', 0)} başlık onay kuyruğundadır."
    )
    rows = []
    for item in criteria:
        measure_rows = item.get("subcriteria", []) or [item]
        for measure in measure_rows:
            main_title = _clean_text(measure.get("main_title"))
            if not _is_criteria_summary_title(main_title):
                continue
            rows.append(
                {
                    "Ana Ölçüt": _report_main_heading(main_title),
                    "Başlık": measure.get("total", 0),
                    "Hazır": measure.get("ready", 0),
                    "Onaylı": measure.get("approved", 0),
                    "Hazırlık %": measure.get("readiness_percent", 0),
                    "Kalite": measure.get("quality_avg", 0),
                }
            )
    _add_table(doc, rows, "Ana Ölçüt Bazlı Değerlendirme Özeti")


def _add_section_content(
    doc: Document,
    username: str,
    program_id: str,
    section: dict[str, Any],
    evidence_targets: dict[str, str],
    base_url: str,
    heading_level: int = 2,
) -> None:
    doc.add_heading(_section_heading_text(section), level=heading_level)
    if _clean_text(section.get("report_text")):
        _add_paragraph_with_evidence_links(doc, _clean_text(section.get("report_text")), evidence_targets)
    else:
        _add_empty_note(doc)

    _add_puko_summary(doc, section, heading_level + 1)

    section_key = _clean_text(section.get("section_key"))
    tables = list_tables(username, program_id, section_key)
    for table in tables:
        _add_table(doc, table.get("rows", []), _clean_text(table.get("table_name")), table.get("meta", {}))
    _add_required_table_markers(doc, section_key, tables)

    evs = list_evidence(username, program_id, section_key)
    if evs:
        doc.add_heading("Kanıtlar", level=min(heading_level + 1, 3))
        for ev in evs:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(f"{ev.get('code')}: ").bold = True
            _add_hyperlink(paragraph, str(ev.get("original_name", "")), _evidence_target(ev, program_id, base_url))
            if ev.get("note"):
                paragraph.add_run(f" - {ev.get('note')}")


def build_final_docx(username: str, program_id: str, base_url: str = "") -> bytes:
    settings = get_settings(program_id)
    sections = list_sections(username, program_id)
    evidence_rows = list_evidence(username, program_id)
    evidence_targets = {str(ev.get("code", "")).lower(): _evidence_target(ev, program_id, base_url) for ev in evidence_rows}

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    _apply_doc_styles(doc)
    _apply_header_footer(doc, settings)

    _add_template_cover(doc, settings)
    _add_contents(doc, sections, _clean_text(settings.get("report_type")) or "ÖZ DEĞERLENDİRME RAPORU")

    doc.add_heading(_clean_text(settings.get("report_type")) or "ÖZ DEĞERLENDİRME RAPORU", level=1)
    for index, (group_title, group_sections) in enumerate(_group_report_sections(sections)):
        if index:
            doc.add_page_break()
        doc.add_heading(group_title, level=1)
        if _is_criteria_group_title(group_title):
            _add_report_summary(doc, username, program_id)
            for subgroup, subgroup_sections in _group_subsections(group_sections):
                doc.add_heading(subgroup, level=2)
                for sec in subgroup_sections:
                    _add_section_content(doc, username, program_id, sec, evidence_targets, base_url, heading_level=3)
        else:
            for sec in group_sections:
                _add_section_content(doc, username, program_id, sec, evidence_targets, base_url, heading_level=2)

    doc.add_page_break()
    doc.add_heading("RAPOR SONU KANIT DİZİNİ", level=1)
    directory_rows = []
    for ev in evidence_rows:
        directory_rows.append(
            {
                "Kanıt Kodu": ev.get("code", ""),
                "Kanıt Adı": ev.get("original_name", ""),
                "İlgili Başlıklar": "; ".join(evidence_links(program_id, str(ev.get("id", "")))),
                "Açıklama": ev.get("note", ""),
                "Yüklenme": ev.get("uploaded_at", ""),
            }
        )
    _add_table(doc, directory_rows, "Konsolide Kanıt Listesi")

    stream = BytesIO()
    doc.save(stream)
    record_export(username, program_id, f"Tam {settings.get('report_short', 'ÖDR')} DOCX", settings.get("docx_filename", "AKYS_ODR.docx"), "FastAPI export")
    return stream.getvalue()


def build_control_docx(username: str, program_id: str) -> bytes:
    settings = get_settings(program_id)
    rows = control_rows(username, program_id)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    _apply_doc_styles(doc)
    _apply_header_footer(doc, settings)
    doc.add_heading(f"{settings.get('accreditation_label', 'MEDEK')} Onay ve Revizyon Kontrol Tablosu", 0)
    doc.add_paragraph(f"Program: {settings.get('program', '')} - {settings.get('report_year', '')}")
    summary = [
        {"Gösterge": "Onaylanan", "Sayı": sum(1 for row in rows if row.get("Onay Durumu") == "Onaylandı")},
        {"Gösterge": "Onay Bekleyen", "Sayı": sum(1 for row in rows if row.get("Onay Durumu") == "Onaya Gönderildi")},
        {"Gösterge": "Revizyon İstenen", "Sayı": sum(1 for row in rows if row.get("Onay Durumu") == "Revizyon Gerekli")},
        {"Gösterge": "Taslak/Hazırlık", "Sayı": sum(1 for row in rows if row.get("Onay Durumu") == "Taslak")},
    ]
    _add_table(doc, summary, "Özet")
    _add_table(doc, rows, "Bölüm Bazlı Kontrol Tablosu")
    stream = BytesIO()
    doc.save(stream)
    record_export(username, program_id, "Kontrol DOCX", settings.get("control_filename", "AKYS_kontrol_tablosu.docx"), "FastAPI export")
    return stream.getvalue()


def build_readiness_audit_docx(username: str, program_id: str) -> bytes:
    settings = get_settings(program_id)
    payload = stats_payload(username, program_id)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    _apply_doc_styles(doc)
    _apply_header_footer(doc, settings)
    doc.add_heading(f"{settings.get('accreditation_label', 'MEDEK')} {settings.get('report_short', 'ÖDR')} Hazırlık Denetimi", 0)
    doc.add_paragraph(f"Program: {settings.get('program', '')} - {settings.get('report_year', '')}")

    summary = payload.get("summary", {})
    totals = payload.get("totals", {})
    _add_table(
        doc,
        [
            {"Gösterge": "Hazırlık Skoru", "Değer": f"{summary.get('readiness_percent', 0)}%"},
            {"Gösterge": "Onay Oranı", "Değer": f"{summary.get('approval_percent', 0)}%"},
            {"Gösterge": "Ortalama Kalite", "Değer": f"{totals.get('avg_quality', 0)}%"},
            {"Gösterge": "Kritik Başlık", "Değer": totals.get("critical_sections", 0)},
            {"Gösterge": "Kanıt", "Değer": totals.get("evidence", 0)},
            {"Gösterge": "Tablo", "Değer": totals.get("tables", 0)},
        ],
        "Denetim Özeti",
    )
    report_group_rows = [
        {
            "Rapor Bölümü": row.get("main_title", ""),
            "Başlık": row.get("total", 0),
            "Hazır": row.get("ready", 0),
            "Onaylı": row.get("approved", 0),
            "Revizyon": row.get("revision", 0),
            "Hazırlık %": row.get("readiness_percent", 0),
            "Kalite": row.get("quality_avg", 0),
        }
        for row in payload.get("report_groups", [])
    ]
    measure_rows = [
        {
            "Ana Ölçüt": row.get("main_title", ""),
            "Başlık": row.get("total", 0),
            "Hazır": row.get("ready", 0),
            "Onaylı": row.get("approved", 0),
            "Revizyon": row.get("revision", 0),
            "Hazırlık %": row.get("readiness_percent", 0),
            "Kalite": row.get("quality_avg", 0),
        }
        for row in payload.get("measure_criteria", payload.get("criteria", []))
    ]
    critical_rows = [
        {
            "Kod": row.get("section_key", ""),
            "Başlık": row.get("section_title", ""),
            "Rapor Bölümü": row.get("report_group_title", ""),
            "Ana Ölçüt": row.get("report_subgroup_title") or row.get("main_title", ""),
            "Kalite": row.get("quality", 0),
            "Kelime": row.get("words", 0),
            "Kanıt": row.get("evidence", 0),
            "Tablo": row.get("tables", 0),
            "PUKÖ": row.get("puko", 0),
        }
        for row in payload.get("critical", [])
    ]
    _add_table(doc, report_group_rows, "Rapor Bölümleri Hazırlık Haritası")
    _add_table(doc, measure_rows, "Ana Ölçüt Bazlı Kalite Haritası")
    _add_table(doc, critical_rows, "Öncelikli Eksik ve Riskli Başlıklar")
    stream = BytesIO()
    doc.save(stream)
    record_export(username, program_id, "Hazırlık Denetimi DOCX", settings.get("audit_filename", "AKYS_hazirlik_denetimi.docx"), "FastAPI export")
    return stream.getvalue()



def build_compliance_audit_docx(username: str, program_id: str) -> bytes:
    """Build a governance/compliance audit package from immutable trail tables."""
    from .enterprise.audit import compliance_audit_payload

    settings = get_settings(program_id)
    payload = compliance_audit_payload(username, program_id, limit=800)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    _apply_doc_styles(doc)
    _apply_header_footer(doc, settings)
    doc.add_heading(f"{settings.get('accreditation_label', 'MEDEK')} Compliance Denetim Raporu", 0)
    doc.add_paragraph(f"Program: {settings.get('program', '')} - {settings.get('report_year', '')}")
    doc.add_paragraph("Bu rapor activity log, onay/revizyon geçmişi, bildirim, çıktı ve bölüm versiyon kayıtlarını denetim amacıyla birleştirir.")

    summary = payload.get("summary", {})
    _add_table(doc, [{"Gösterge": key, "Değer": value} for key, value in summary.items()], "Compliance Özet")
    _add_table(doc, payload.get("approval_counts", []), "Onay Durumu Dağılımı")
    _add_table(doc, payload.get("actor_counts", [])[:30], "Kullanıcı / Aktör Dağılımı")
    _add_table(doc, payload.get("action_counts", [])[:40], "İşlem Türü Dağılımı")
    _add_table(doc, payload.get("section_activity", [])[:80], "Başlık Bazlı Hareket Yoğunluğu")
    _add_table(doc, payload.get("stale_workflow", [])[:80], "Bekleyen Workflow Öğeleri")
    _add_table(doc, payload.get("risk_rows", [])[:40], "Riskli Başlıklar")
    _add_table(doc, payload.get("activity", [])[:200], "Son Activity Log Kayıtları")
    _add_table(doc, payload.get("approvals", [])[:200], "Onay / Revizyon Geçmişi")
    _add_table(doc, payload.get("notifications", [])[:200], "Bildirim Geçmişi")
    _add_table(doc, payload.get("exports", [])[:100], "Çıktı Geçmişi")
    _add_table(doc, payload.get("versions", [])[:200], "Versiyon Snapshot Geçmişi")

    stream = BytesIO()
    doc.save(stream)
    record_export(username, program_id, "Compliance Audit DOCX", "AKYS_compliance_audit.docx", "FastAPI export")
    return stream.getvalue()



def _bar_text(value: float, max_value: float = 100.0, width: int = 18) -> str:
    try:
        numeric = float(value or 0)
    except (TypeError, ValueError):
        numeric = 0.0
    ratio = 0.0 if max_value <= 0 else max(0.0, min(1.0, numeric / max_value))
    filled = int(round(ratio * width))
    return "█" * filled + "░" * max(0, width - filled)


def build_advanced_analytics_docx(username: str, program_id: str) -> bytes:
    """Build an executive analytics dashboard export as DOCX."""
    from .enterprise.dashboard import advanced_reporting

    settings = get_settings(program_id)
    payload = advanced_reporting(username, program_id)
    summary = payload.get("summary", {})
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    _apply_doc_styles(doc)
    _apply_header_footer(doc, settings)
    doc.add_heading(f"{settings.get('accreditation_label', 'MEDEK')} Advanced Analytics Dashboard", 0)
    doc.add_paragraph(f"Program: {settings.get('program', '')} - {settings.get('report_year', '')}")
    doc.add_paragraph("Bu yönetici özeti; tamamlanma oranı, onay yoğunluğu, PUKÖ doluluğu, risk ısı haritası ve işlem trendlerini denetim/yönetim amacıyla raporlar.")

    _add_table(
        doc,
        [
            {"Gösterge": "Toplam Başlık", "Değer": summary.get("total", 0)},
            {"Gösterge": "Onaylanan Başlık", "Değer": summary.get("approved", 0)},
            {"Gösterge": "Revizyon Bekleyen", "Değer": summary.get("revision", 0)},
            {"Gösterge": "Yüksek Riskli Başlık", "Değer": summary.get("high_risk", 0)},
            {"Gösterge": "Ortalama Kalite", "Değer": f"{summary.get('quality_avg', 0)}%"},
        ],
        "Yönetici KPI Özeti",
    )

    group_rows = []
    for row in payload.get("group_chart", []):
        group_rows.append({
            "Rapor Grubu": row.get("group", ""),
            "Başlık": row.get("total", 0),
            "Hazır": row.get("ready", 0),
            "Onaylı": row.get("approved", 0),
            "Revizyon": row.get("revision", 0),
            "Hazırlık %": row.get("readiness_percent", 0),
            "Onay %": row.get("approval_percent", 0),
            "Kalite": row.get("quality_avg", 0),
            "Hazırlık Bar": _bar_text(float(row.get("readiness_percent", 0) or 0)),
        })
    _add_table(doc, group_rows, "Grup Bazlı Tamamlanma ve Onay Grafiği")

    total = max(1, int(summary.get("total", 0) or 0))
    puko_rows = []
    for row in payload.get("puko_chart", []):
        count = int(row.get("count", 0) or 0)
        percent = round(count / total * 100, 1)
        puko_rows.append({"PUKÖ Alanı": row.get("field", ""), "Dolu Başlık": count, "Doluluğu %": percent, "Bar": _bar_text(percent)})
    _add_table(doc, puko_rows, "PUKÖ Doluluk Dağılımı")

    _add_table(doc, payload.get("approval_distribution", []), "Onay Durumu Dağılımı")
    _add_table(doc, payload.get("status_distribution", []), "Hazırlık Durumu Dağılımı")

    trend_rows = []
    for row in payload.get("trend_chart", []):
        trend_rows.append({"Tarih": row.get("date", ""), "Kaydedilen Bölüm Snapshot": row.get("saved_sections", 0)})
    _add_table(doc, trend_rows, "Bölüm Versiyon Trendleri")

    activity_rows = []
    for row in payload.get("activity_trend", []):
        activity_rows.append({"Tarih": row.get("date", ""), "Aktivite Sayısı": row.get("activity_count", 0)})
    _add_table(doc, activity_rows, "Kullanıcı Aktivite Trendleri")

    risk_rows = []
    for row in payload.get("risk_heatmap", [])[:80]:
        risk = int(row.get("risk", 0) or 0)
        level = "Kritik" if risk >= 70 else "Orta" if risk >= 45 else "Düşük"
        risk_rows.append({
            "Kod": row.get("section_key", ""),
            "Başlık": row.get("section_title", ""),
            "Grup": row.get("group", ""),
            "Risk": risk,
            "Seviye": level,
            "Kalite": row.get("quality", 0),
            "Durum": row.get("status", ""),
            "Onay": row.get("approval_status", ""),
        })
    _add_table(doc, risk_rows, "Risk Heat Map - Öncelikli Başlıklar")

    doc.add_heading("Yönetici Yorumu", level=1)
    high_risk = int(summary.get("high_risk", 0) or 0)
    revision = int(summary.get("revision", 0) or 0)
    quality_avg = float(summary.get("quality_avg", 0) or 0)
    recommendations = []
    if high_risk:
        recommendations.append(f"{high_risk} yüksek riskli başlık için kanıt, tablo ve PUKÖ metinleri öncelikli tamamlanmalıdır.")
    if revision:
        recommendations.append(f"{revision} revizyon bekleyen başlık için sorumlu editör/onaylayıcı takibi yapılmalıdır.")
    if quality_avg < 70:
        recommendations.append("Ortalama kalite %70 altında olduğu için metin derinliği, kanıt bağlantısı ve PUKÖ bütünlüğü güçlendirilmelidir.")
    if not recommendations:
        recommendations.append("Genel hazırlık görünümü kabul edilebilir düzeydedir; onaylanmamış başlıklar ve güncel kanıtlar düzenli izlenmelidir.")
    for item in recommendations:
        doc.add_paragraph(item, style=None)

    stream = BytesIO()
    doc.save(stream)
    record_export(username, program_id, "Advanced Analytics DOCX", "AKYS_advanced_analytics_dashboard.docx", "FastAPI export")
    return stream.getvalue()

def convert_docx_to_pdf(docx_data: bytes, base_name: str) -> bytes:
    candidates = [
        os.environ.get("MEDEK_SOFFICE_PATH", ""),
        shutil.which("soffice") or "",
        shutil.which("libreoffice") or "",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        docx_path = tmp_path / f"{slugify(base_name) or 'AKYS_ODR'}.docx"
        pdf_path = docx_path.with_suffix(".pdf")
        docx_path.write_bytes(docx_data)
        for candidate in candidates:
            if candidate and Path(candidate).exists():
                subprocess.run(
                    [candidate, "--headless", "--convert-to", "pdf", "--outdir", str(tmp_path), str(docx_path)],
                    capture_output=True,
                    text=True,
                    timeout=120,
                    check=False,
                )
                if pdf_path.exists() and pdf_path.stat().st_size > 0:
                    return pdf_path.read_bytes()
        if os.name == "nt":
            try:
                subprocess.run(
                    [
                        sys.executable,
                        "-c",
                        "from docx2pdf import convert; import sys; convert(sys.argv[1], sys.argv[2])",
                        str(docx_path),
                        str(pdf_path),
                    ],
                    capture_output=True,
                    text=True,
                    timeout=120,
                    check=False,
                )
                if pdf_path.exists() and pdf_path.stat().st_size > 0:
                    return pdf_path.read_bytes()
            except Exception:
                pass
    raise RuntimeError("PDF üretimi için LibreOffice/soffice veya Windows'ta Microsoft Word + docx2pdf gerekir.")
